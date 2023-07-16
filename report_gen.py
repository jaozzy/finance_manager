from docx import Document
from manager import extra_infos
from datetime import datetime
from essencials import get_path, file_exists

# Get the path of the script
path = get_path()
graphdir = f'{path}/graphs/'

# Check if the file 'relatorio.docx' exists
file_exists('relatorio.docx')

def generate_report():
    # Gets the exact moment when the report was generated
    agora = datetime.now()
    data_formatada = agora.strftime("%d/%m/%Y")
    
    # Import the necessary informations from manager.py
    receitas, receitas_total, receitas_met, receitas_cat, despesas, despesas_total, despesas_met, despesas_cat, supv, percD_R, percD_M, percD_C, percR_M, percR_C, data_min, data_max, periodo, periodo_str, mdia_d_D, mdia_r_D = extra_infos()
    
    doc = Document()

    doc.add_heading('Relatório Financeiro', level=1)  # Add a level 1 heading for the financial report
    doc.add_paragraph(f'Relatório Gerado em {data_formatada} às {agora.strftime("%H:%M")}')  # Add the report generation date and time
    doc.add_paragraph('Desenvolvido por: João Pedro')  # Add developer name
    doc.add_paragraph('Contato: (47) 9 99783190 / servicecontact.joao@gmail.com')  # Add developer contact information

    doc.add_heading('Informações Gerais:', level=2)  # Add a level 2 heading for the general information section
    doc.add_paragraph(f'Total gasto no período analisado: R${despesas_total}')  # Add total expenses in the analyzed period
    doc.add_paragraph(f'Total recebido no período analisado: R${receitas_total}')  # Add total revenues in the analyzed period
    doc.add_paragraph(f'Superávit/Défict: R${supv}')  # Add surplus/deficit amount
    doc.add_paragraph(f'Data mais antiga analisada: {data_min.strftime("%d/%m/%Y")}')  # Add earliest analyzed date
    doc.add_paragraph(f'Data mais recente analisada: {data_max.strftime("%d/%m/%Y")}')  # Add latest analyzed date
    doc.add_paragraph(f'Período de Análise: {periodo_str}')  # Add analysis period

    doc.add_heading('Detalhes de Gastos por Categoria', level=2)  # Add a level 2 heading for expense details by category

    # Loop through expense categories and add details to the document
    for despesa, categoria in despesas_cat.items():
        despesa, categoria = categoria, despesa
        mdia = despesa / periodo.days if periodo.days != 0 else 0
        gperc = (despesa / despesas_total) * 100
        doc.add_paragraph(f'O total gasto em {categoria} no período analisado foi de {despesa}R$')
        doc.add_paragraph(f'A média de gasto diário com {categoria} foi de aproximadamente {mdia:.2f}R$, representando {gperc:.2f}% do valor total de gastos do período analisado.')
    doc.add_paragraph()
        
    doc.add_heading('Gráfico de Setores Relacionado às Despesas por Categoria em Função das Despesas totais do Periodo Analisado:', level=2)  # Add a level 2 heading for the pie chart of expenses by category
    doc.add_picture(f'{graphdir}/grafico_despesas_categoria.png')  # Add the pie chart image for expenses by category
    
    doc.add_page_break()
    
    doc.add_heading('Detalhes de Gastos por Método de Pagamento', level=2)  # Add a level 2 heading for expense details by payment method

    # Loop through expense payment methods and add details to the document
    for despesa, metodo in despesas_met.items():
        despesa, metodo = metodo, despesa
        mdia = (despesa / periodo.days)
        gperc = (despesa / despesas_total) * 100
        doc.add_paragraph(f'O total gasto por {metodo} no período analisado foi de {despesa}R$')
        doc.add_paragraph(f'A média de gasto diário por {metodo} foi de aproximadamente {mdia:.2f}R$, representando {gperc:.2f}% do valor total de gastos do período analisado.')
        doc.add_paragraph()
        
    doc.add_heading('Gráfico de Setores Relacionado às Despesas por Método de Pagamento em Função das Despesas totais do Periodo Analisado:', level=2)  # Add a level 2 heading for the pie chart of expenses by payment method
    doc.add_picture(f'{graphdir}/grafico_despesas_metodo.png')  # Add the pie chart image for expenses by payment method
    
    doc.add_page_break()

    doc.add_heading('Detalhes de Receitas por Categoria:', level=2)  # Add a level 2 heading for revenue details by category

    # Loop through revenue categories and add details to the document
    for receita, categoria in receitas_cat.items():
        receita, categoria = categoria, receita
        mdia = (receita / periodo.days)
        rperc = (receita / receitas_total) * 100
        doc.add_paragraph(f'O total recebido em {categoria} no período analisado foi de R${receita}')
        doc.add_paragraph(f'A média de recebimento diário foi de aproximadamente R${mdia:.2f}, representando {rperc:.2f}% do valor total de gastos do período analisado.')
        doc.add_paragraph()
        
    doc.add_heading('Gráfico de Setores Relacionado às Receitas por Categoria em Função das Receitas totais do Periodo Analisado:', level=2)  # Add a level 2 heading for the pie chart of revenues by category
    doc.add_picture(f'{graphdir}/grafico_receitas_categoria.png')  # Add the pie chart image for revenues by category
    
    doc.add_page_break()
    
    doc.add_heading('Detalhes de Receitas por Método de Pagamento:', level=2)  # Add a level 2 heading for revenue details by payment method

    # Loop through revenue payment methods and add details to the document
    for receita, metodo in receitas_met.items():
        receita, metodo = metodo, receita
        mdia = (receita / periodo.days)
        rperc = (receita / receitas_total) * 100
        doc.add_paragraph(f'O total recebido por {metodo} no período analisado foi de R${receita}')
        doc.add_paragraph(f'A média de recebimento diário foi de aproximadamente R${mdia:.2f}, representando {rperc:.2f}% do valor total de gastos do período analisado.')
        doc.add_paragraph()
        
    doc.add_heading('Gráfico de Setores Relacionado às Receitas por Método de Pagamento em Função das Receitas totais do Periodo Analisado:', level=2)  # Add a level 2 heading for the pie chart of revenues by payment method
    doc.add_picture(f'{graphdir}/grafico_receitas_metodo.png')  # Add the pie chart image for revenues by payment method
    
    doc.save('relatorio.docx')

generate_report()